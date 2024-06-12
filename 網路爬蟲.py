import requests
from bs4 import BeautifulSoup
from docx import Document
import tkinter as tk
from tkinter import messagebox
from tkinter import filedialog

def fetch_page_content(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        return response.text
    except requests.exceptions.HTTPError as http_err:
        print(f"HTTP錯誤: {http_err}")
    except Exception as err:
        print(f"其他錯誤: {err}")

def parse_html(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    
    title = soup.title.string if soup.title else '無標題'
    texts = list(soup.stripped_strings)
    links = [a['href'] for a in soup.find_all('a', href=True)]
    images = [img['src'] for img in soup.find_all('img', src=True)]
    
    return title, texts, links, images

def save_to_word(title, texts, links, images, filename):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_heading('文本內容', level=1)
    for text in texts:
        doc.add_paragraph(text)
    doc.add_heading('連結', level=1)
    for link in links:
        doc.add_paragraph(link)
    doc.add_heading('圖片', level=1)
    for image in images:
        doc.add_paragraph(image)
    doc.save(filename)

def start_scraping():
    url = url_entry.get()
    
    if not url:
        messagebox.showwarning("輸入錯誤", "請輸入URL。")
        return
    
    filename = filedialog.asksaveasfilename(defaultextension=".docx",
                                            filetypes=[("Word文件", "*.docx"), ("所有文件", "*.*")])
    
    if not filename:
        messagebox.showwarning("輸入錯誤", "請選擇文件名和路徑。")
        return
    
    html_content = fetch_page_content(url)
    if not html_content:
        messagebox.showerror("獲取錯誤", "無法獲取網頁內容。")
        return
    
    title, texts, links, images = parse_html(html_content)
    save_to_word(title, texts, links, images, filename)
    messagebox.showinfo("成功", f"內容已保存到 {filename}")

# 創建圖形化介面
root = tk.Tk()
root.title("網頁爬蟲")

tk.Label(root, text="請輸入URL:").grid(row=0, column=0, padx=10, pady=10)
url_entry = tk.Entry(root, width=50)
url_entry.grid(row=0, column=1, padx=10, pady=10)

tk.Button(root, text="選擇文件保存位置並開始爬取", command=start_scraping).grid(row=2, column=0, columnspan=2, pady=20)

root.mainloop()
